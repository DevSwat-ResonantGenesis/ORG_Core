"""
Document Loaders for RAG Pipeline (Phase 1.2)
==============================================

Parse uploaded documents into text for chunking + embedding.
Supports: PDF, DOCX, CSV, HTML, TXT, MD, JSON

Each loader returns a list of text chunks ready for embedding.
"""

import csv
import io
import json
import logging
import re
from typing import List, Optional

logger = logging.getLogger(__name__)


def parse_document(content: bytes, filename: str, content_type: Optional[str] = None) -> str:
    """Parse a document into plain text based on file extension / content type.
    
    Returns extracted text or raises ValueError if unsupported.
    """
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()

    # PDF
    if ext == "pdf" or (content_type and "pdf" in content_type):
        return _parse_pdf(content)

    # DOCX
    if ext == "docx" or (content_type and "wordprocessingml" in (content_type or "")):
        return _parse_docx(content)

    # CSV
    if ext == "csv" or (content_type and "csv" in (content_type or "")):
        return _parse_csv(content)

    # HTML
    if ext in ("html", "htm") or (content_type and "html" in (content_type or "")):
        return _parse_html(content)

    # JSON
    if ext == "json" or (content_type and "json" in (content_type or "")):
        return _parse_json(content)

    # Plain text / markdown fallback
    return _decode_text(content)


def chunk_text(text: str, max_chars: int = 1000, overlap: int = 100) -> List[str]:
    """Split text into overlapping chunks by paragraph boundaries.
    
    Tries to split on double-newlines first, then single newlines, then force-split.
    """
    if not text or not text.strip():
        return []

    paragraphs = re.split(r"\n\s*\n", text)
    chunks: List[str] = []
    current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        if len(current) + len(para) + 2 <= max_chars:
            current = (current + "\n\n" + para) if current else para
        else:
            if current.strip():
                chunks.append(current.strip())
            # If single paragraph exceeds max, split on sentences
            if len(para) > max_chars:
                sentences = re.split(r"(?<=[.!?])\s+", para)
                current = ""
                for sent in sentences:
                    if len(current) + len(sent) + 1 <= max_chars:
                        current = (current + " " + sent) if current else sent
                    else:
                        if current.strip():
                            chunks.append(current.strip())
                        # Force-split very long sentences
                        if len(sent) > max_chars:
                            for i in range(0, len(sent), max_chars - overlap):
                                chunks.append(sent[i:i + max_chars].strip())
                            current = ""
                        else:
                            current = sent
            else:
                current = para

    if current.strip():
        chunks.append(current.strip())

    # Add overlap between chunks
    if overlap > 0 and len(chunks) > 1:
        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            prev_tail = chunks[i - 1][-overlap:] if len(chunks[i - 1]) > overlap else chunks[i - 1]
            overlapped.append(prev_tail + "\n" + chunks[i])
        chunks = overlapped

    return [c for c in chunks if len(c) >= 10]


# ---------------------------------------------------------------------------
# Individual parsers
# ---------------------------------------------------------------------------

def _decode_text(content: bytes) -> str:
    """Decode bytes to text with fallback encodings."""
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return content.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return content.decode("utf-8", errors="replace")


def _parse_pdf(content: bytes) -> str:
    """Extract text from PDF using PyPDF2 or pdfplumber."""
    # Try PyPDF2 first (lighter)
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        if pages:
            return "\n\n".join(pages)
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"PyPDF2 failed: {e}")

    # Try pdfplumber (better extraction)
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        if pages:
            return "\n\n".join(pages)
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")

    raise ValueError("PDF parsing requires PyPDF2 or pdfplumber. Install: pip install PyPDF2")


def _parse_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs)
    except ImportError:
        raise ValueError("DOCX parsing requires python-docx. Install: pip install python-docx")
    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {e}")


def _parse_csv(content: bytes) -> str:
    """Convert CSV to readable text (header: value format)."""
    text = _decode_text(content)
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)

    if not rows:
        return ""

    headers = rows[0] if rows else []
    lines = []

    if len(rows) > 1 and headers:
        for row in rows[1:]:
            parts = []
            for i, val in enumerate(row):
                h = headers[i] if i < len(headers) else f"col{i}"
                if val.strip():
                    parts.append(f"{h}: {val}")
            if parts:
                lines.append("; ".join(parts))
    else:
        for row in rows:
            lines.append(", ".join(row))

    return "\n".join(lines)


def _parse_html(content: bytes) -> str:
    """Strip HTML tags and extract text."""
    text = _decode_text(content)

    # Try BeautifulSoup
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(text, "html.parser")
        # Remove script/style
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except ImportError:
        pass

    # Fallback: regex strip
    text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _parse_json(content: bytes) -> str:
    """Convert JSON to readable text."""
    text = _decode_text(content)
    try:
        data = json.loads(text)
        return json.dumps(data, indent=2, ensure_ascii=False, default=str)
    except json.JSONDecodeError:
        return text
